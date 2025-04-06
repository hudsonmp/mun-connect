import io
import os
import logging
from typing import Dict, Optional

logger = logging.getLogger(__name__)

def extract_text_from_file(file_data: bytes, file_type: str) -> Optional[str]:
    """
    Extract text content from various file types.
    
    Args:
        file_data: Binary file data
        file_type: MIME type of file
        
    Returns:
        Extracted text or None if extraction fails
    """
    try:
        if file_type == 'application/pdf':
            return extract_from_pdf(file_data)
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            return extract_from_docx(file_data)
        elif file_type == 'text/plain':
            return file_data.decode('utf-8', errors='replace')
        else:
            logger.warning(f"Unsupported file type for extraction: {file_type}")
            return None
    except Exception as e:
        logger.error(f"Error extracting text from file: {str(e)}")
        return None

def extract_from_pdf(file_data: bytes) -> str:
    """Extract text from a PDF file"""
    text = ""
    try:
        # Lazy import to avoid importing unless needed
        import PyPDF2
        
        pdf_file = io.BytesIO(file_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Simple extraction with text sanitization
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            
            # Basic text cleanup
            if page_text:
                # Replace multiple spaces and newlines
                cleaned_text = ' '.join(page_text.split())
                text += cleaned_text + "\n\n"
        
        return text.strip()
    except ImportError:
        logger.error("PyPDF2 not installed. Cannot extract text from PDF.")
        return "Error: PDF parsing library not available. Could not extract text."
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        # Return partial text if available
        return text if text else "Error extracting text from PDF"

def extract_from_docx(file_data: bytes) -> str:
    """Extract text from a DOCX file"""
    text = ""
    try:
        # Lazy import to avoid importing unless needed
        import docx
        
        doc_file = io.BytesIO(file_data)
        doc = docx.Document(doc_file)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except ImportError:
        logger.error("python-docx not installed. Cannot extract text from DOCX.")
        return "Error: DOCX parsing library not available. Could not extract text."
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        # Return partial text if available
        return text if text else "Error extracting text from DOCX" 